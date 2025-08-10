# rag_pipeline.py
# This file contains the core logic for the Retrieval-Augmented Generation pipeline
# with table extraction, reranking, caching, and parallel processing.

import os
import re
import requests
import faiss
import torch
import networkx as nx
import nltk
import numpy as np
import hashlib
import pdfplumber
import docx
import eml_parser
from sklearn.preprocessing import normalize
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer, CrossEncoder
from PyPDF2 import PdfReader
from openai import OpenAI
from pydantic import BaseModel
from typing import List
import logging
from concurrent.futures import ThreadPoolExecutor

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

nltk.data.path.append('/tmp/nltk_data')

nltk.download("punkt_tab", quiet=True)

class RAGPipeline:
    def __init__(self, openai_api_key: str, 
                 embed_model_name: str = "BAAI/bge-small-en-v1.5",
                 rerank_model_name: str = "cross-encoder/ms-marco-MiniLM-L-6-v2"):
        """
        Initializes the RAG pipeline, loading models and setting up configurations.
        """
        logger.info("Initializing RAG Pipeline with Caching and Reranker...")
        
        # --- Configuration ---
        self.CHUNK_SIZE = 300
        self.OVERLAP = 50
        self.TOP_K_RETRIEVAL = 12
        self.FINAL_K_RERANK = 5
        self.CACHE_DIR = "cache"

        # --- Model and API Setup ---
        if not openai_api_key:
            raise ValueError("OPENAI_API_KEY is required.")
        
        self.openai_client = OpenAI(api_key=openai_api_key)
        
        device = 'cuda' if torch.cuda.is_available() else 'cpu'
        logger.info(f"Loading models onto device: {device}...")
        self.embedder = SentenceTransformer(embed_model_name, device=device)
        self.reranker = CrossEncoder(rerank_model_name, device=device)
        logger.info("All models loaded.")

        # --- State variables ---
        self.chunks = None
        self.faiss_index = None
        self.graph = None
        self.is_ready = False
        os.makedirs(self.CACHE_DIR, exist_ok=True)
        logger.info("RAG Pipeline initialized.")

    def _hash_url(self, url: str) -> str:
        return hashlib.sha256(url.encode("utf-8")).hexdigest()

    def _save_cache(self, prefix: str, chunks: list, embeddings: np.ndarray, faiss_index):
        logger.info(f"Saving cache for prefix: {prefix}")
        np.save(os.path.join(self.CACHE_DIR, f"{prefix}_embeddings.npy"), embeddings)
        faiss.write_index(faiss_index, os.path.join(self.CACHE_DIR, f"{prefix}_faiss.index"))
        with open(os.path.join(self.CACHE_DIR, f"{prefix}_chunks.txt"), "w", encoding="utf-8") as f:
            f.write("\n".join(chunks))

    def _load_cache(self, prefix: str):
        logger.info(f"Attempting to load cache for prefix: {prefix}")
        embedding_path = os.path.join(self.CACHE_DIR, f"{prefix}_embeddings.npy")
        faiss_path = os.path.join(self.CACHE_DIR, f"{prefix}_faiss.index")
        chunks_path = os.path.join(self.CACHE_DIR, f"{prefix}_chunks.txt")

        if not (os.path.exists(embedding_path) and os.path.exists(faiss_path) and os.path.exists(chunks_path)):
            logger.info("Cache not found.")
            return None, None, None
        
        embeddings = np.load(embedding_path)
        faiss_index = faiss.read_index(faiss_path)
        with open(chunks_path, "r", encoding="utf-8") as f:
            chunks = f.read().splitlines()
        
        logger.info("Cache loaded successfully.")
        return chunks, embeddings, faiss_index

    def _table_to_markdown(self, table_data):
        if not table_data:
            return ""
        num_cols = max(len(row) for row in table_data)
        markdown_string = ""
        header = [(cell if cell is not None else "") for cell in table_data[0]]
        header += [""] * (num_cols - len(header))
        markdown_string += "| " + " | ".join(str(cell).replace("|", "\\|") for cell in header) + " |\n"
        markdown_string += "|" + "|".join(["---"] * num_cols) + "|\n"
        for row in table_data[1:]:
            row = [(cell if cell is not None else "") for cell in row]
            row += [""] * (num_cols - len(row))
            markdown_string += "| " + " | ".join(str(cell).replace("|", "\\|") for cell in row) + " |\n"
        return markdown_string

    def _download_and_extract_text(self, url: str) -> str:
        """
        Downloads a file from a URL and extracts text content based on its type.
        """
        logger.info(f"Downloading and extracting text from {url}...")
        
        # Determine file extension from URL
        file_extension = url.split('?')[0].split('.')[-1].lower()
        local_path = f"downloaded_doc.{file_extension}"

        with open(local_path, "wb") as f:
            f.write(requests.get(url).content)

        full_text = ""
        if file_extension == 'pdf':
            with pdfplumber.open(local_path) as pdf:
                # (Your existing PDF extraction logic)
                for page in pdf.pages:
                    full_text += page.extract_text() + "\n"
                    for table in page.find_tables():
                        if table.extract():
                            full_text += "\n" + self._table_to_markdown(table.extract()) + "\n"
        
        elif file_extension == 'docx':
            doc = docx.Document(local_path)
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                full_text += "\n" + self._table_to_markdown(table_data) + "\n"

        elif file_extension == 'eml':
            with open(local_path, 'rb') as f:
                raw_email = f.read()
            ep = eml_parser.EmlParser()
            parsed_eml = ep.decode_email_bytes(raw_email)
            if parsed_eml['body']:
                full_text = parsed_eml['body'][0]['content']

        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
        return full_text

    def _chunk_text(self, text: str) -> list[str]:
        logger.info("Performing smart chunking...")
        sentences = nltk.sent_tokenize(text)
        chunks, current_chunk_words, current_length = [], [], 0
        for sentence in sentences:
            words = re.findall(r"\S+", sentence)
            if current_length + len(words) > self.CHUNK_SIZE and current_chunk_words:
                chunks.append(" ".join(current_chunk_words))
                overlap_word_count = int(self.OVERLAP * len(current_chunk_words) / current_length) if current_length > 0 else 0
                current_chunk_words = current_chunk_words[-overlap_word_count:]
                current_length = len(current_chunk_words)
            current_chunk_words.extend(words)
            current_length += len(words)
        if current_chunk_words:
            chunks.append(" ".join(current_chunk_words))
        return chunks

    def _get_embeddings(self, chunks: list[str]) -> np.ndarray:
        return normalize(self.embedder.encode([f"passage: {c}" for c in chunks], convert_to_numpy=True, batch_size=32))

    def _build_faiss(self, embeddings: np.ndarray):
        index = faiss.IndexFlatIP(embeddings.shape[1])
        index.add(embeddings)
        return index

    def _build_graph(self, chunks: list[str], embeddings: np.ndarray):
        G = nx.Graph()
        for i, text in enumerate(chunks):
            G.add_node(i, text=text)
        sim_matrix = cosine_similarity(embeddings)
        for i in range(len(chunks)):
            for j in range(i + 1, len(chunks)):
                if sim_matrix[i][j] > 0.75:
                    G.add_edge(i, j, weight=sim_matrix[i][j])
        return G

    def process_document(self, doc_url: str):
        cache_prefix = self._hash_url(doc_url)
        cached_chunks, cached_embeddings, cached_faiss = self._load_cache(cache_prefix)
        
        if cached_chunks and cached_faiss is not None:
            self.chunks = cached_chunks
            self.faiss_index = cached_faiss
            self.graph = self._build_graph(self.chunks, cached_embeddings)
        else:
            logger.info("Cache not found, processing document from scratch...")
            raw_text = self._download_and_extract_text(doc_url)
            self.chunks = self._chunk_text(raw_text)
            embeddings = self._get_embeddings(self.chunks)
            self.faiss_index = self._build_faiss(embeddings)
            self.graph = self._build_graph(self.chunks, embeddings)
            self._save_cache(cache_prefix, self.chunks, embeddings, self.faiss_index)

        self.is_ready = True
        logger.info("--- Document processing complete. Pipeline is ready. ---")

    def _retrieve_chunks(self, query: str) -> list[str]:
        if not self.is_ready:
            raise RuntimeError("Pipeline not ready.")
        q_embed = normalize(self.embedder.encode([f"query: {query}"], convert_to_numpy=True))
        _, indices = self.faiss_index.search(q_embed, self.TOP_K_RETRIEVAL)
        
        base_indices = set(indices[0])
        for i in indices[0]:
            if self.graph.has_node(i) and list(self.graph.neighbors(i)):
                 base_indices.update(list(self.graph.neighbors(i))[:3])
        
        return [self.chunks[i] for i in sorted(list(base_indices))]

    def _rerank_chunks(self, query: str, candidates: list[str]) -> list[str]:
        pairs = [(query, passage) for passage in candidates]
        scores = self.reranker.predict(pairs, batch_size=32)
        ranked = sorted(zip(candidates, scores), key=lambda x: x[1], reverse=True)
        return [text for text, _ in ranked[:self.FINAL_K_RERANK]]

    def _generate_answer(self, original_query: str, top_chunks: list[str]) -> str:
        context = "\n\n".join(top_chunks)
        prompt = f"""You are answering a question using the following context from an insurance policy. 
Don't use knowledge which is not in the context. Keep the answer brief.

Context:
{context}

Question:
{original_query}

Answer:"""
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Error during OpenAI API call: {e}")
            return "Error: Could not generate an answer."

    def _process_single_question(self, question: str) -> str:
        """Helper function to process one question for parallel execution."""
        retrieved = self._retrieve_chunks(question)
        reranked = self._rerank_chunks(question, retrieved)
        answer = self._generate_answer(question, reranked)
        logger.info(f"Question-----> {question}")
        logger.info(f"Answer-------> {answer}")
        return answer

    def answer_questions(self, questions: list[str]) -> list[str]:
        """
        Answers a list of questions in parallel and returns only the answer strings.
        """
        if not self.is_ready:
            raise RuntimeError("Pipeline not ready.")
        
        logger.info(".........ANSWERING QUESTIONS........")
        
        with ThreadPoolExecutor() as executor:
            answers = list(executor.map(self._process_single_question, questions))
            
        return answers